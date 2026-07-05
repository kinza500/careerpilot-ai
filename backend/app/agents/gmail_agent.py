"""Gmail integration: drafting applications/follow-ups, checking threads for
replies, and (with explicit per-event user confirmation) creating calendar
events when a reply turns out to mention a scheduled interview.

Scopes: `gmail.compose` (create drafts), `gmail.readonly` (read message
content, needed to tell whether a reply arrived and what it says), and
`calendar.events` (create/edit calendar events only — cannot read the rest of
the user's calendar). None of these can send mail or auto-confirm a calendar
event on their own; sending stays a manual Gmail action, and calendar events
are only ever created after the user clicks to confirm one, consistent with
the project's human-in-the-loop design throughout. gmail.readonly technically
grants read access to the whole mailbox, not just application threads (Gmail
has no narrower "one thread" scope); this code only ever reads the specific
thread_id stored against an application.
"""
from __future__ import annotations

import base64
import io
from email import encoders
from email.mime.application import MIMEApplication
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.config import get_settings

settings = get_settings()

AUTH_URL = "https://accounts.google.com/o/oauth2/v2/auth"
TOKEN_URL = "https://oauth2.googleapis.com/token"
USERINFO_URL = "https://www.googleapis.com/oauth2/v2/userinfo"
DRAFTS_URL = "https://gmail.googleapis.com/gmail/v1/users/me/drafts"
MESSAGES_URL = "https://gmail.googleapis.com/gmail/v1/users/me/messages"
THREADS_URL = "https://gmail.googleapis.com/gmail/v1/users/me/threads"
CALENDAR_EVENTS_URL = "https://www.googleapis.com/calendar/v3/calendars/primary/events"

SCOPES = (
    "https://www.googleapis.com/auth/gmail.compose "
    "https://www.googleapis.com/auth/gmail.readonly "
    "https://www.googleapis.com/auth/calendar.events "
    "openid email"
)


class GoogleTokenRevoked(Exception):
    """Google has confirmed the refresh token itself is dead (invalid_grant —
    actually revoked or expired), as opposed to a transient network/API
    failure. Callers should only treat THIS as "disconnect the user" — never
    a generic exception, which could just as easily be a dropped connection."""


def build_authorize_url(state: str) -> str:
    params = {
        "response_type": "code",
        "client_id": settings.google_client_id,
        "redirect_uri": settings.google_redirect_uri,
        "scope": SCOPES,
        "access_type": "offline",
        "prompt": "consent",
        "state": state,
    }
    return f"{AUTH_URL}?{httpx.QueryParams(params)}"


async def exchange_code(code: str) -> dict:
    async with httpx.AsyncClient() as client:
        res = await client.post(TOKEN_URL, data={
            "code": code,
            "client_id": settings.google_client_id,
            "client_secret": settings.google_client_secret,
            "redirect_uri": settings.google_redirect_uri,
            "grant_type": "authorization_code",
        })
        res.raise_for_status()
        return res.json()


async def refresh_access_token(refresh_token: str) -> str:
    async with httpx.AsyncClient() as client:
        res = await client.post(TOKEN_URL, data={
            "refresh_token": refresh_token,
            "client_id": settings.google_client_id,
            "client_secret": settings.google_client_secret,
            "grant_type": "refresh_token",
        })
        if res.status_code == 400:
            try:
                error = res.json().get("error")
            except ValueError:
                error = None
            if error == "invalid_grant":
                raise GoogleTokenRevoked("Google refresh token is invalid or has been revoked")
        res.raise_for_status()
        return res.json()["access_token"]


async def get_email(access_token: str) -> str | None:
    async with httpx.AsyncClient() as client:
        res = await client.get(USERINFO_URL, headers={"Authorization": f"Bearer {access_token}"})
        if res.status_code != 200:
            return None
        return res.json().get("email")


TOKENINFO_URL = "https://oauth2.googleapis.com/tokeninfo"


async def has_readonly_scope(access_token: str) -> bool | None:
    """Whether the CURRENTLY GRANTED token actually includes gmail.readonly —
    not just what SCOPES asks for. A token issued before this scope existed
    in the app (or where the user's Google consent didn't include it) will
    silently 403 on every thread check otherwise, with reply/sent-detection
    looking like it's just not working rather than a permissions gap.
    Returns None if the check itself fails (treat as unknown, not missing)."""
    async with httpx.AsyncClient() as client:
        try:
            res = await client.get(TOKENINFO_URL, params={"access_token": access_token})
        except httpx.HTTPError:
            return None
    if res.status_code != 200:
        return None
    granted = res.json().get("scope", "").split()
    return "https://www.googleapis.com/auth/gmail.readonly" in granted


async def has_calendar_scope(access_token: str) -> bool | None:
    """Same check as has_readonly_scope, for the calendar.events scope —
    added after gmail.readonly, so an account connected before this feature
    existed will need to reconnect to grant it."""
    async with httpx.AsyncClient() as client:
        try:
            res = await client.get(TOKENINFO_URL, params={"access_token": access_token})
        except httpx.HTTPError:
            return None
    if res.status_code != 200:
        return None
    granted = res.json().get("scope", "").split()
    return "https://www.googleapis.com/auth/calendar.events" in granted


DOCX_MIME = "vnd.openxmlformats-officedocument.wordprocessingml.document"


def _text_to_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in (text or "").split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _resume_to_docx_bytes(text: str) -> bytes:
    """Renders the tailored resume into a properly laid-out Word document —
    name + contact centred at top, bold section headings, real bullet lists
    — instead of dumping the LLM's plain text as one flat wall of paragraphs.
    Relies on RESUME_OPTIMISATION's output contract (name on line 1, contact
    on line 2, ALL-CAPS section headings, "- " bullets); anything that
    doesn't match those patterns just falls back to a plain paragraph, so
    unexpected formatting degrades gracefully rather than breaking."""
    doc = Document()
    lines = [l.rstrip() for l in (text or "").split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)

    if lines:
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_p.add_run(lines.pop(0).strip())
        run.bold = True
        run.font.size = Pt(18)

    if lines and lines[0].strip():
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.add_run(lines.pop(0).strip()).font.size = Pt(10)

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("- ", "• ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped.isupper() and len(stripped.split()) <= 4:
            doc.add_heading(stripped, level=2)
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def create_draft(
    access_token: str, subject: str, body: str, to: str | None = None,
    cover_letter: str | None = None, tailored_resume: str | None = None,
    original_resume: tuple[bytes, str, str] | None = None,
    thread_id: str | None = None, in_reply_to: str | None = None,
) -> dict:
    """Create a Gmail draft. Passing thread_id + in_reply_to (the original
    message's RFC822 Message-ID header) threads this into an existing
    conversation instead of starting a new one — used for follow-up emails.

    original_resume, if given, is (raw_bytes, filename, mime_type) for the
    EXACT file the user uploaded — attached completely unchanged (a PDF stays
    a PDF), so its real layout is guaranteed rather than risking a lossy
    reformat/conversion. This is separate from tailored_resume, the
    LLM-rewritten version formatted into its own Word doc.
    """
    msg = MIMEMultipart()
    msg["Subject"] = subject
    if to:
        msg["To"] = to
    if in_reply_to:
        msg["In-Reply-To"] = in_reply_to
        msg["References"] = in_reply_to
    msg.attach(MIMEText(body))

    attachments = []
    if cover_letter:
        attachments.append(("Cover Letter.docx", cover_letter, _text_to_docx_bytes))
    if tailored_resume:
        attachments.append(("Tailored Resume.docx", tailored_resume, _resume_to_docx_bytes))
    for filename, text, renderer in attachments:
        part = MIMEApplication(renderer(text), _subtype=DOCX_MIME)
        part.add_header("Content-Disposition", "attachment", filename=filename)
        msg.attach(part)

    if original_resume:
        data, filename, mime_type = original_resume
        maintype, _, subtype = (mime_type or "application/octet-stream").partition("/")
        part = MIMEBase(maintype or "application", subtype or "octet-stream")
        part.set_payload(data)
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", "attachment", filename=filename)
        msg.attach(part)

    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    message: dict = {"raw": raw}
    if thread_id:
        message["threadId"] = thread_id

    async with httpx.AsyncClient() as client:
        res = await client.post(
            DRAFTS_URL,
            headers={"Authorization": f"Bearer {access_token}"},
            json={"message": message},
        )
        if res.status_code >= 400:
            print(f"[gmail] draft create failed {res.status_code}: {res.text}")
        res.raise_for_status()
        return res.json()


async def get_message_id_header(access_token: str, message_id: str) -> str | None:
    """Fetch the RFC822 Message-ID header of a message (needed to thread a
    follow-up via In-Reply-To/References)."""
    async with httpx.AsyncClient() as client:
        res = await client.get(
            f"{MESSAGES_URL}/{message_id}",
            headers={"Authorization": f"Bearer {access_token}"},
            params={"format": "metadata", "metadataHeaders": "Message-ID"},
        )
        if res.status_code != 200:
            return None
        for h in res.json().get("payload", {}).get("headers", []):
            if h.get("name", "").lower() == "message-id":
                return h.get("value")
        return None


async def thread_has_reply(access_token: str, thread_id: str) -> bool | None:
    """Whether a reply has arrived on an application's thread.

    A message lands with the INBOX label only when Gmail actually receives
    it from someone else — our own unsent draft only ever carries the DRAFT
    label, and our own sent message carries SENT, so INBOX presence is an
    unambiguous "someone replied" signal without needing to read any message
    body.

    Returns None (not False) when the check can't be performed at all — a
    revoked/insufficient-scope token, a deleted thread, or a network error —
    so callers fall back to the time-based nudge instead of assuming no reply.
    """
    async with httpx.AsyncClient() as client:
        try:
            res = await client.get(
                f"{THREADS_URL}/{thread_id}",
                headers={"Authorization": f"Bearer {access_token}"},
                params={"format": "metadata"},
            )
        except httpx.HTTPError:
            return None
    if res.status_code != 200:
        return None
    messages = res.json().get("messages", [])
    return any("INBOX" in m.get("labelIds", []) for m in messages)


async def count_sent_in_thread(access_token: str, thread_id: str) -> int | None:
    """How many messages in a thread carry the SENT label — i.e. messages the
    user actually sent themselves from Gmail, as opposed to unsent drafts
    (DRAFT label only). Used to detect "was this draft actually sent" by
    comparing against a baseline count captured when the draft was created —
    a plain boolean "any SENT message" isn't enough once a thread can contain
    both an original application email and a later follow-up.

    Returns None (not 0) when the check can't be performed at all — a
    revoked/insufficient-scope token, a deleted thread, or a network error —
    so callers treat it as "unknown" rather than "definitely 0".
    """
    async with httpx.AsyncClient() as client:
        try:
            res = await client.get(
                f"{THREADS_URL}/{thread_id}",
                headers={"Authorization": f"Bearer {access_token}"},
                params={"format": "metadata"},
            )
        except httpx.HTTPError:
            return None
    if res.status_code != 200:
        return None
    messages = res.json().get("messages", [])
    return sum(1 for m in messages if "SENT" in m.get("labelIds", []))


def _decode_part(data: str | None) -> str:
    if not data:
        return ""
    try:
        return base64.urlsafe_b64decode(data + "=" * (-len(data) % 4)).decode("utf-8", errors="replace")
    except Exception:
        return ""


def _extract_body_text(payload: dict) -> str:
    """Walk a Gmail message payload for the best available text — prefers
    text/plain, falls back to a crude tag-strip of text/html."""
    mime = payload.get("mimeType", "")
    if mime == "text/plain":
        return _decode_part(payload.get("body", {}).get("data"))
    if mime == "text/html":
        import re
        html = _decode_part(payload.get("body", {}).get("data"))
        return re.sub(r"<[^>]+>", " ", html)
    for part in payload.get("parts", []) or []:
        text = _extract_body_text(part)
        if text.strip():
            return text
    return _decode_part(payload.get("body", {}).get("data"))


async def get_latest_reply_text(access_token: str, thread_id: str) -> str | None:
    """Full plain-text content of the most recent INBOX-labelled (i.e.
    genuinely received, not our own draft/sent) message in a thread — used
    to check whether a reply mentions a scheduled interview. Only ever reads
    the one thread already stored against an application, never browses the
    mailbox generally."""
    async with httpx.AsyncClient() as client:
        try:
            res = await client.get(
                f"{THREADS_URL}/{thread_id}",
                headers={"Authorization": f"Bearer {access_token}"},
                params={"format": "full"},
            )
        except httpx.HTTPError:
            return None
    if res.status_code != 200:
        return None
    messages = [m for m in res.json().get("messages", []) if "INBOX" in m.get("labelIds", [])]
    if not messages:
        return None
    latest = messages[-1]  # Gmail returns thread messages in chronological order
    return _extract_body_text(latest.get("payload", {}))[:6000]


async def create_calendar_event(
    access_token: str, summary: str, description: str,
    start_iso: str, end_iso: str, timezone: str = "UTC", location: str | None = None,
) -> dict:
    """Create an event on the user's primary calendar — only ever called
    after the user explicitly confirms a suggested interview slot; nothing
    here runs automatically off the back of just reading an email."""
    body = {
        "summary": summary,
        "description": description,
        "start": {"dateTime": start_iso, "timeZone": timezone},
        "end": {"dateTime": end_iso, "timeZone": timezone},
    }
    if location:
        body["location"] = location
    async with httpx.AsyncClient() as client:
        res = await client.post(
            CALENDAR_EVENTS_URL,
            headers={"Authorization": f"Bearer {access_token}"},
            json=body,
        )
        if res.status_code >= 400:
            print(f"[gmail] calendar event create failed {res.status_code}: {res.text}")
        res.raise_for_status()
        return res.json()
